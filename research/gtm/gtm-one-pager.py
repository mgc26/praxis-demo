#!/usr/bin/env python3
"""Generate GTM one-pager DOCX for Praxis Demo Platform."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page setup: narrow margins for one-pager density --
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

TEAL = RGBColor(0x00, 0xB9, 0xCE)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
    h.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    h.paragraph_format.space_after = Pt(4)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = GRAY
    return p

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)

# ============================================================================
# TITLE
# ============================================================================
title = doc.add_heading('Praxis AI Agent Platform', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title.runs:
    run.font.color.rgb = DARK
    run.font.size = Pt(22)
    run.font.name = 'Calibri'
title.paragraph_format.space_after = Pt(2)

subtitle = doc.add_paragraph('Demo Sales Guide  |  Life Sciences & Specialty Pharma')
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in subtitle.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL
    run.font.name = 'Calibri'
    run.bold = True
subtitle.paragraph_format.space_after = Pt(6)

# Divider line
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run('_' * 95)
run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
run.font.size = Pt(6)

# ============================================================================
# ELEVATOR PITCH
# ============================================================================
add_heading_styled('What This Demo Shows', level=2)
add_body(
    'An AI-native patient and HCP engagement platform purpose-built for specialty pharma. '
    'Four specialized agents handle the full lifecycle of pharmaceutical support\u2014from '
    'patient hub enrollment and adherence coaching to HCP medical information and field '
    'engagement\u2014with regulatory compliance, adverse event capture, and audit trails '
    'built into every interaction.'
)

# ============================================================================
# JTBD TABLE
# ============================================================================
add_heading_styled('Jobs-to-Be-Done (4 Agents, 4 Workflows)', level=2)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Agent', 'Direction', 'Primary Jobs', 'Key Differentiator']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

rows_data = [
    [
        'Patient Support\n(Emma)',
        'Outbound',
        'Hub enrollment, copay activation, adherence coaching, AE capture, nurse educator scheduling',
        'Caregiver support & break-in-therapy detection for rare disease',
    ],
    [
        'HCP Support\n(Aria)',
        'Inbound',
        'Medical information inquiries, dosing questions, drug interactions, AE/pregnancy reporting, MSL referral',
        'On-label boundary enforcement with automatic off-label routing to MSL',
    ],
    [
        'HCP Outbound\n(Marcus)',
        'Outbound',
        'Signal-driven detail calls, competitive differentiation, formulary barrier resolution, speaker program capture',
        'Behavioral signal triggers (competitor research, formulary lookups, conference activity)',
    ],
    [
        'MedComms QA\n(Rachel)',
        'Internal',
        'Promotional compliance review, off-label deviation detection, AE gap audits, crisis protocol documentation',
        'Real-time compliance QA before content reaches the field',
    ],
]

for r, row_data in enumerate(rows_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r + 1].cells[c]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                if c == 0:
                    run.bold = True

# ============================================================================
# DEMO FEATURES TO HIGHLIGHT
# ============================================================================
add_heading_styled('Recommended Demo Flow & Features to Call Out', level=2)

add_bullet(
    ' Start here. Walk through a single interaction from behavioral signal detection '
    'to agent configuration to live conversation to transcript analysis to call log entry. '
    'This is the "aha" moment\u2014show how an AI agent replaces a traditional call center rep '
    'with pharma-grade intelligence.',
    bold_prefix='Agent Storyboard \u2192',
)
add_bullet(
    ' Show the prioritized contact queue with behavioral signals '
    '(competitor research detected, adherence gap, formulary lookup). Each contact arrives '
    'pre-scored with a recommended pathway and priority tier\u2014no manual triage needed.',
    bold_prefix='Interaction Data \u2192',
)
add_bullet(
    ' MSL follow-up requests ranked by urgency, AE severity flags, '
    'competitive intelligence capture with corroboration counts. This is what the medical '
    'affairs team sees the morning after 50 agent interactions.',
    bold_prefix='Medical Intelligence \u2192',
)
add_bullet(
    ' Engagement funnel (signals \u2192 outreach \u2192 contact \u2192 action), '
    'outcome distribution across all four agent types, daily trend lines. '
    'Connects agent activity to commercial KPIs.',
    bold_prefix='Performance Analytics \u2192',
)
add_bullet(
    ' Patient cohort tracking, timepoint-based outcome data (baseline through 6-month), '
    'payer evidence cards for formulary discussions. Demonstrates real-world evidence generation '
    'from agent interactions.',
    bold_prefix='Outcomes & Evidence \u2192',
)

# ============================================================================
# COMPLIANCE AS A FEATURE
# ============================================================================
add_heading_styled('Compliance Built In, Not Bolted On', level=2)

add_bullet('FDA adverse event capture with structured minimum elements (onset, severity, causality, reporter)')
add_bullet('Anti-kickback screening before copay card activation (federal program beneficiary exclusion)')
add_bullet('On-label boundary enforcement with documented off-label routing to MSL')
add_bullet('Pregnancy exposure registry auto-enrollment for anti-epileptic drugs')
add_bullet('PhRMA Code / Sunshine Act compliance for HCP outbound engagement')
add_bullet('Full audit trail: every interaction classified, timestamped, and available for regulatory review')

# ============================================================================
# MULTI-BRAND
# ============================================================================
add_heading_styled('Multi-Brand Architecture', level=2)

add_body(
    'The platform supports multiple brand packs from a single deployment. The demo includes '
    'three pre-configured brands to show enterprise portfolio readiness:'
)

add_bullet(
    ' ELEX (Essential Tremor) + Relutrigine (Dravet Syndrome)\u2014specialty neurology with rare disease complexity',
    bold_prefix='Praxis Precision Medicines:',
)
add_bullet(
    ' Repatha (ASCVD) + TEZSPIRE (Severe Asthma)\u2014large-molecule cardio-metabolic and respiratory',
    bold_prefix='Amgen:',
)
add_bullet(
    ' Sephience (PKU) + Emflaza (DMD)\u2014rare disease with genetic testing integration',
    bold_prefix='PTC Therapeutics:',
)
add_body(
    'Each brand pack carries its own drug profiles, therapeutic areas, support pathways, '
    'agent personas, visual identity, and compliance guardrails. New brands can be onboarded '
    'without re-architecting the platform.'
)

# ============================================================================
# WHO CARES
# ============================================================================
add_heading_styled('Who This Resonates With', level=2)

table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

persona_rows = [
    ['Stakeholder', 'What They Care About'],
    ['Patient Services / Hub Ops', 'Faster enrollment, copay activation, adherence lift, reduced call center headcount'],
    ['Medical Affairs / MSLs', 'Pre-call intelligence, follow-up request triage, competitive intel capture, off-label routing'],
    ['Commercial / Brand Ops', 'Signal-driven outreach ROI, prescriber engagement lift, formulary access support'],
    ['Pharmacovigilance', 'First-touch AE capture, pregnancy exposure reporting, 15-day expedited reporting visibility'],
    ['Compliance / Legal', 'Real-time promotional deviation detection, audit trail completeness, guardrail enforcement'],
]

for r, (col1, col2) in enumerate(persona_rows):
    table2.rows[r].cells[0].text = col1
    table2.rows[r].cells[1].text = col2
    for c in range(2):
        for p in table2.rows[r].cells[c].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                if r == 0 or c == 0:
                    run.bold = True

# ============================================================================
# TALK TRACK TIPS
# ============================================================================
add_heading_styled('Talk Track Tips', level=2)

add_bullet(
    ' "What if your hub services could detect a caregiver in distress from their search '
    'behavior\u2014before they ever pick up the phone\u2014and proactively reach out with the right support?"',
    bold_prefix='Lead with the signal: ',
)
add_bullet(
    ' "Every agent captures adverse events with FDA minimum elements, screens for '
    'anti-kickback before activating copay cards, and enforces on-label boundaries automatically. '
    'Compliance isn\'t a gate\u2014it\'s embedded."',
    bold_prefix='Compliance as value: ',
)
add_bullet(
    ' "This isn\'t a single-brand tool. We demo three brands\u2014neurology, cardio-metabolic, '
    'rare disease\u2014from the same platform. Your entire portfolio, one deployment."',
    bold_prefix='Multi-brand proof: ',
)
add_bullet(
    ' "After 50 agent interactions, your MSL team has a prioritized follow-up queue, '
    'your PV team has structured AE cases, and your commercial team has engagement funnel data\u2014'
    'all without a single manual note."',
    bold_prefix='Downstream value: ',
)

# ============================================================================
# FOOTER
# ============================================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('_' * 95)
run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
run.font.size = Pt(6)

footer = doc.add_paragraph('Internal Use Only  |  Vi Technologies  |  March 2026')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'
    run.italic = True

# Save
out_path = os.path.join(os.path.dirname(__file__), 'Praxis-Demo-GTM-One-Pager.docx')
doc.save(out_path)
print(f'Saved to {out_path}')
